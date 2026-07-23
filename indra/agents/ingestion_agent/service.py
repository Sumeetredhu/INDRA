"""End-to-end ingestion service.

The modules in :mod:`indra.agents.ingestion_agent` deliberately keep format parsing, OCR,
chunking, tag normalisation, and P&ID vision independently testable.  This service is the only
place where those stages become one durable workflow and where a parsed document is handed to the
knowledge-graph service through its public contract.
"""

from __future__ import annotations

import asyncio
import csv
import io
import re
import time
from collections.abc import Sequence
from datetime import date
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import TYPE_CHECKING

from indra.agents.ingestion_agent.chunking import SemanticChunker, TextBlock
from indra.agents.ingestion_agent.ocr import OCREngine
from indra.agents.ingestion_agent.pid_vision import PIDVisionParser, pid_to_graph
from indra.agents.ingestion_agent.tag_normalizer import PlantTagNormalizer, equipment_type_for
from indra.agents.ingestion_agent.validation import ValidationReport, read_and_validate, validate_bytes
from indra.core.contracts import KnowledgeGraphService, ProgressCallback
from indra.core.deps import AgentDeps
from indra.core.events import Event, Topic
from indra.core.exceptions import IndraError, IngestionError
from indra.core.ids import content_id, new_id
from indra.core.logging import get_logger
from indra.core.models import (
    Confidence,
    DocumentMeta,
    EntityType,
    ExtractedEntity,
    ExtractedRelationship,
    IngestionProgress,
    IngestionResult,
    IngestionStage,
    MimeFamily,
    ParsedDocument,
    RelationType,
)

if TYPE_CHECKING:  # pragma: no cover - imported only for static analysis
    pass

logger = get_logger(__name__)

_TAG_RE = re.compile(r"(?<![A-Z0-9])(?:[A-Z]{1,3})[-\s]?[0-9OILSBZ]{2,4}[A-Z]?(?![A-Z0-9])", re.IGNORECASE)
_DATE_RE = re.compile(r"\b(20\d{2})[-_](\d{2})[-_](\d{2})\b")
_FAILURE_RE = re.compile(
    r"\b(bearing\s+(?:seizure|wear|failure)|lubrication\s+failure|oil\s+pump\s+malfunction|"
    r"cavitation|overheating|vibration)\b",
    re.IGNORECASE,
)
_PERSON_RE = re.compile(r"\b([A-Z][a-z]+\s+[A-Z][a-z]+)\b")
_SECTION_RE = re.compile(r"\b(?:section|clause|rule)\s+([\dA-Za-z().-]+)", re.IGNORECASE)


class IngestionAgent:
    """Validate, store, parse, enrich, and index any supported plant document."""

    name = "ingestion_agent"

    def __init__(self, deps: AgentDeps) -> None:
        self._deps = deps
        self._chunker = SemanticChunker(deps.settings)
        self._ocr = OCREngine(deps.settings)
        self._normalizer = PlantTagNormalizer(deps.settings)
        self._pid = PIDVisionParser(deps.settings, ocr=self._ocr, normalizer=self._normalizer)
        self._knowledge_graph: KnowledgeGraphService | None = None

    def bind(self, *, knowledge_graph: KnowledgeGraphService) -> None:
        """Attach the public graph-service boundary after all agents are constructed."""
        self._knowledge_graph = knowledge_graph

    async def startup(self) -> None:
        """Warm the tag registry without making startup depend on a healthy graph."""
        try:
            equipment = await self._deps.graph.list_equipment()
            self._normalizer.set_registry([item.tag for item in equipment])
        except Exception as exc:
            logger.warning("could not warm ingestion tag registry", extra={"error": str(exc)})

    async def shutdown(self) -> None:
        """The service owns no closeable resources; lifecycle hook satisfies the wiring contract."""

    async def health(self) -> dict[str, object]:
        """Return an honest capability summary for the operations panel."""
        ocr = await self._ocr.describe()
        return {
            "ok": True,
            "backend": "pipeline",
            "detail": f"chunker={self._chunker.name}, pid_detector={self._pid.detector_name}",
            "ocr": ocr,
            "tag_registry_size": len(self._normalizer.registry),
        }

    async def ingest_path(
        self, path: Path, *, on_progress: ProgressCallback | None = None
    ) -> IngestionResult:
        """Read a local file on a worker thread, then run the byte workflow."""
        try:
            content, _ = await asyncio.to_thread(read_and_validate, path, settings=self._deps.settings)
        except IndraError:
            raise
        except Exception as exc:  # pragma: no cover - defensive filesystem boundary
            raise IngestionError(
                f"Could not read {path.name} for ingestion. Check that the file is available.",
                context={"path": str(path)},
                cause=exc,
            ) from exc
        return await self.ingest_bytes(content, filename=path.name, on_progress=on_progress)

    async def ingest_directory(self, directory: Path, *, concurrency: int = 4) -> list[IngestionResult]:
        """Ingest a flat corpus with bounded concurrency and deterministic result ordering."""
        if not directory.is_dir():
            raise IngestionError(
                f"{directory} is not a readable ingestion directory.", context={"directory": str(directory)}
            )
        files = sorted(path for path in directory.iterdir() if path.is_file())
        semaphore = asyncio.Semaphore(max(1, concurrency))

        async def one(path: Path) -> IngestionResult:
            async with semaphore:
                return await self.ingest_path(path)

        return list(await asyncio.gather(*(one(path) for path in files)))

    async def ingest_bytes(
        self,
        content: bytes,
        *,
        filename: str,
        on_progress: ProgressCallback | None = None,
    ) -> IngestionResult:
        """Run the complete ingestion pipeline for a single upload.

        Duplicate bytes short-circuit before parsing, retaining the original document identifier as
        required by content-addressed ingestion.  All other pipeline boundaries emit progress and
        typed events so a client can render a real-time pipeline rather than a fake progress bar.
        """
        started = time.perf_counter()
        job_id = new_id("job")
        report = await asyncio.to_thread(validate_bytes, content, filename=filename, settings=self._deps.settings)
        document_id = content_id(content, kind="document")
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.RECEIVED, 2.0, "Upload received")
        await self._publish(Topic.DOCUMENT_RECEIVED, job_id=job_id, document_id=document_id, filename=report.filename)
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.VALIDATED, 10.0, "Validated magic number and size")

        duplicate = await self._deps.metadata.find_by_hash(report.content_hash)
        if duplicate is not None and self._deps.settings.ingestion_idempotent:
            result = IngestionResult(
                job_id=job_id,
                document=duplicate,
                duplicate_of=duplicate.document_id,
                warnings=["Identical content already ingested; reused the existing document."],
                duration_ms=round((time.perf_counter() - started) * 1000.0, 2),
            )
            await self._progress(on_progress, job_id, duplicate.document_id, report.filename, IngestionStage.COMPLETE, 100.0, "Duplicate content reused")
            await self._publish(Topic.DOCUMENT_INGESTED, job_id=job_id, document_id=duplicate.document_id, duplicate_of=duplicate.document_id, chunks=0, entities=0, relationships=0)
            return result

        uri = await self._deps.blobs.put(content, filename=report.filename, content_hash=report.content_hash)
        meta = self._meta(report, document_id=document_id, source_path=uri)
        await self._deps.metadata.save_document(meta)
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.STORED, 20.0, "Stored raw document")

        registry = await self._equipment_tags()
        blocks, tables, warnings = await self._extract(content, report, meta, registry=registry)
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.PARSED, 38.0, "Extracted document content")

        chunks, text = await self._chunker.chunk(blocks, document_id=document_id)
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.CHUNKED, 52.0, f"Created {len(chunks)} semantic chunks", chunks=len(chunks))

        entities, relationships = self._extract_entities(chunks, registry=registry)
        pid_result = None
        if report.mime_family is MimeFamily.IMAGE and meta.document_type.value == "pid_drawing":
            try:
                local_path = await self._deps.blobs.path_for(uri)
                pid_result = await self._pid.parse(local_path, document_id=document_id, registry=registry)
                vision_entities, vision_relationships = pid_to_graph(pid_result, meta)
                entities.extend(vision_entities)
                relationships.extend(vision_relationships)
                warnings.extend(pid_result.warnings)
            except Exception as exc:
                warnings.append(f"P&ID vision degraded: {type(exc).__name__}: {exc}")
                logger.warning("P&ID vision degraded", extra={"document_id": document_id, "error": str(exc)})

        parsed = ParsedDocument(
            meta=meta,
            text=text,
            chunks=chunks,
            entities=entities,
            relationships=relationships,
            tables=tables,
            pid_result=pid_result,
            warnings=[*report.warnings, *warnings],
            stage=IngestionStage.GRAPH_QUEUED,
            parse_duration_ms=round((time.perf_counter() - started) * 1000.0, 2),
        )
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.ENTITIES_EXTRACTED, 67.0, f"Found {len(entities)} entities", entities=len(entities), relationships=len(relationships), chunks=len(chunks))
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.RELATIONS_EXTRACTED, 75.0, f"Inferred {len(relationships)} relationships", entities=len(entities), relationships=len(relationships), chunks=len(chunks))

        if self._knowledge_graph is None:
            raise IngestionError("Ingestion agent is not bound to the knowledge graph service. Start through IndraOrchestrator.")
        counts = await self._knowledge_graph.index(parsed)
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.GRAPH_WRITTEN, 92.0, "Indexed graph and vector evidence", entities=len(entities), relationships=len(relationships), chunks=len(chunks))

        result = IngestionResult(
            job_id=job_id,
            document=meta,
            chunks_created=counts.get("chunks", len(chunks)),
            entities_created=counts.get("entities", len(entities)),
            relationships_created=counts.get("relationships", len(relationships)),
            pid_symbols=len(pid_result.symbols) if pid_result else 0,
            pid_connections=len(pid_result.connections) if pid_result else 0,
            warnings=parsed.warnings,
            duration_ms=round((time.perf_counter() - started) * 1000.0, 2),
        )
        await self._progress(on_progress, job_id, document_id, report.filename, IngestionStage.COMPLETE, 100.0, "Ingestion complete", entities=result.entities_created, relationships=result.relationships_created, chunks=result.chunks_created)
        await self._publish(
            Topic.DOCUMENT_INGESTED,
            job_id=job_id,
            document_id=document_id,
            title=meta.title,
            document_type=meta.document_type.value,
            chunks=result.chunks_created,
            entities=result.entities_created,
            relationships=result.relationships_created,
            is_pid=pid_result is not None,
        )
        return result

    def _meta(self, report: ValidationReport, *, document_id: str, source_path: str) -> DocumentMeta:
        """Build metadata without trusting client-provided type declarations."""
        title = Path(report.filename).stem.replace("_", " ").strip() or report.filename
        dated = _DATE_RE.search(report.filename)
        document_date: date | None = None
        if dated:
            try:
                document_date = date(int(dated.group(1)), int(dated.group(2)), int(dated.group(3)))
            except ValueError:
                document_date = None
        return DocumentMeta(
            document_id=document_id,
            title=title,
            filename=report.filename,
            content_hash=report.content_hash,
            mime_family=report.mime_family,
            mime_type=report.mime_type,
            document_type=report.document_type,
            size_bytes=report.size_bytes,
            document_date=document_date,
            source_path=source_path,
        )

    async def _equipment_tags(self) -> list[str]:
        try:
            return [item.tag for item in await self._deps.graph.list_equipment()]
        except Exception as exc:
            logger.warning("equipment registry unavailable for tag normalisation", extra={"error": str(exc)})
            return list(self._normalizer.registry)

    async def _extract(
        self,
        content: bytes,
        report: ValidationReport,
        meta: DocumentMeta,
        *,
        registry: Sequence[str],
    ) -> tuple[list[TextBlock], list[dict[str, object]], list[str]]:
        """Route content to a safe parser and return provenance-bearing text blocks."""
        try:
            if report.mime_family is MimeFamily.PDF:
                return await asyncio.to_thread(self._pdf_blocks, content)
            if report.mime_family is MimeFamily.WORD:
                return await asyncio.to_thread(self._word_blocks, content)
            if report.mime_family is MimeFamily.SPREADSHEET:
                return await asyncio.to_thread(self._spreadsheet_blocks, content, report.extension)
            if report.mime_family is MimeFamily.EMAIL:
                return await asyncio.to_thread(self._email_blocks, content)
            if report.mime_family is MimeFamily.IMAGE:
                # The OCR protocol accepts decoded images, not transport bytes.  Decode here so
                # an uploaded diagram follows the same path as a local P&ID image.
                from PIL import Image

                image = await asyncio.to_thread(
                    lambda: Image.open(io.BytesIO(content)).convert("RGB").copy()
                )
                result = await self._ocr.recognize(image)
                blocks = [TextBlock(text=result.text, page=1, kind="paragraph", ocr_confidence=result.mean_confidence)] if result.ok else []
                warnings = list(result.warnings)
                if not result.ok:
                    warnings.append("Image OCR returned no text; visual P&ID analysis will still be attempted when applicable.")
                return blocks, [], warnings
            return self._text_blocks(content), [], []
        except IndraError:
            raise
        except Exception as exc:
            raise IngestionError(
                f"Could not parse {report.filename}. The raw file was retained; verify it is not encrypted or corrupt.",
                context={"filename": report.filename, "mime_family": report.mime_family.value},
                cause=exc,
            ) from exc

    @staticmethod
    def _pdf_blocks(content: bytes) -> tuple[list[TextBlock], list[dict[str, object]], list[str]]:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        blocks = [TextBlock(text=(page.extract_text() or "").strip(), page=index, kind="paragraph") for index, page in enumerate(reader.pages, start=1)]
        blocks = [block for block in blocks if block.text]
        warnings = [] if blocks else ["PDF contained no extractable text; a scanned-PDF OCR renderer is not configured."]
        return blocks, [], warnings

    @staticmethod
    def _word_blocks(content: bytes) -> tuple[list[TextBlock], list[dict[str, object]], list[str]]:
        from docx import Document

        document = Document(io.BytesIO(content))
        blocks = [TextBlock(text=paragraph.text.strip(), kind="paragraph") for paragraph in document.paragraphs if paragraph.text.strip()]
        tables: list[dict[str, object]] = []
        for table_index, table in enumerate(document.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            tables.append({"name": f"table_{table_index + 1}", "rows": rows})
            blocks.append(TextBlock(text=" | ".join(" | ".join(row) for row in rows if any(row)), kind="table"))
        return blocks, tables, []

    @staticmethod
    def _spreadsheet_blocks(content: bytes, extension: str) -> tuple[list[TextBlock], list[dict[str, object]], list[str]]:
        tables: list[dict[str, object]] = []
        blocks: list[TextBlock] = []
        if extension.lower() in {".csv", ".tsv"}:
            dialect = csv.excel_tab if extension.lower() == ".tsv" else csv.excel
            text = content.decode("utf-8-sig", errors="replace")
            rows = [[cell.strip() for cell in row] for row in csv.reader(io.StringIO(text), dialect=dialect)]
            tables.append({"name": "data", "rows": rows})
            blocks.extend(TextBlock(text=" | ".join(row), kind="table") for row in rows if any(row))
            return blocks, tables, []
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        try:
            for sheet in workbook.worksheets:
                rows = [["" if value is None else str(value).strip() for value in row] for row in sheet.iter_rows(values_only=True)]
                populated = [row for row in rows if any(row)]
                tables.append({"name": sheet.title, "rows": populated})
                blocks.extend(TextBlock(text=" | ".join(row), section=sheet.title, kind="table") for row in populated)
        finally:
            workbook.close()
        return blocks, tables, []

    @staticmethod
    def _email_blocks(content: bytes) -> tuple[list[TextBlock], list[dict[str, object]], list[str]]:
        message = BytesParser(policy=policy.default).parsebytes(content)
        headers = [f"{field}: {message.get(field, '')}" for field in ("From", "To", "Subject", "Date") if message.get(field)]
        body = message.get_body(preferencelist=("plain",))
        text = body.get_content() if body is not None else message.get_content()
        return [TextBlock(text="\n".join([*headers, str(text)]).strip(), kind="metadata")], [], []

    @staticmethod
    def _text_blocks(content: bytes) -> list[TextBlock]:
        text = content.decode("utf-8", errors="replace")
        return [TextBlock(text=piece.strip(), kind="paragraph") for piece in re.split(r"\n\s*\n", text) if piece.strip()]

    def _extract_entities(
        self, chunks: Sequence[object], *, registry: Sequence[str]
    ) -> tuple[list[ExtractedEntity], list[ExtractedRelationship]]:
        """Perform deterministic plant-domain entity and co-occurrence extraction."""
        entities: list[ExtractedEntity] = []
        relationships: list[ExtractedRelationship] = []
        for raw_chunk in chunks:
            # ``SemanticChunker`` guarantees Chunk instances; this guard keeps the internal helper total.
            chunk = raw_chunk
            text = getattr(chunk, "text", "")
            document_id = getattr(chunk, "document_id", None)
            chunk_id = getattr(chunk, "chunk_id", None)
            page = getattr(chunk, "page", None)
            local: list[ExtractedEntity] = []
            seen_tags: set[str] = set()
            known_surfaces = {re.sub(r"[-\s_]", "", item).upper() for item in registry}
            for match in _TAG_RE.finditer(text):
                surface = match.group(0)
                # A permissive OCR-aware grammar can read ordinary prose such as ``at 90`` as an
                # instrument tag.  Preserve lowercase tags only when they reconcile to a known
                # asset; a newly discovered born-digital tag must retain its conventional capitals.
                flat_surface = re.sub(r"[-\s_]", "", surface).upper()
                if surface != surface.upper() and flat_surface not in known_surfaces:
                    continue
                tag, score, alternatives = self._normalizer.normalize(surface, registry=registry)
                if tag is None or tag in seen_tags:
                    continue
                seen_tags.add(tag)
                local.append(ExtractedEntity(
                    type=EntityType.EQUIPMENT,
                    name=match.group(0),
                    canonical_name=tag,
                    confidence=Confidence(value=score, rationale="Plant-tag grammar and registry reconciliation", method="heuristic"),
                    document_id=document_id,
                    chunk_id=chunk_id,
                    page=page,
                    char_start=match.start(),
                    char_end=match.end(),
                    alternatives=alternatives,
                    attributes={"equipment_type": equipment_type_for(tag)},
                ))
            for match in _FAILURE_RE.finditer(text):
                local.append(ExtractedEntity(
                    type=EntityType.FAILURE_MODE,
                    name=match.group(0),
                    confidence=Confidence(value=0.86, rationale="Plant-domain failure-mode phrase", method="heuristic"),
                    document_id=document_id,
                    chunk_id=chunk_id,
                    page=page,
                    char_start=match.start(),
                    char_end=match.end(),
                ))
            for match in _SECTION_RE.finditer(text):
                local.append(ExtractedEntity(
                    type=EntityType.REGULATORY_CLAUSE,
                    name=f"Section {match.group(1)}",
                    confidence=Confidence(value=0.92, rationale="Regulatory clause pattern", method="heuristic"),
                    document_id=document_id,
                    chunk_id=chunk_id,
                    page=page,
                    char_start=match.start(),
                    char_end=match.end(),
                ))
            for match in _PERSON_RE.finditer(text):
                name = match.group(1)
                if name.lower() in {"factory act", "root cause", "bearing wear"}:
                    continue
                local.append(ExtractedEntity(
                    type=EntityType.PERSON,
                    name=name,
                    confidence=Confidence(value=0.62, rationale="Title-case personnel name pattern", method="heuristic"),
                    document_id=document_id,
                    chunk_id=chunk_id,
                    page=page,
                    char_start=match.start(),
                    char_end=match.end(),
                ))
            entities.extend(local)
            equipment = [entity for entity in local if entity.type is EntityType.EQUIPMENT]
            failures = [entity for entity in local if entity.type is EntityType.FAILURE_MODE]
            for asset in equipment:
                for failure in failures:
                    relationships.append(ExtractedRelationship(
                        type=RelationType.FAILED_WITH_MODE,
                        source_key=asset.key,
                        target_key=failure.key,
                        confidence=Confidence(value=0.72, rationale="Equipment and failure mode co-occur in one passage", method="heuristic"),
                        evidence_text=text[:600],
                        document_id=document_id,
                        chunk_id=chunk_id,
                        method="co_occurrence",
                    ))
        return entities, relationships

    async def _progress(
        self,
        callback: ProgressCallback | None,
        job_id: str,
        document_id: str,
        filename: str,
        stage: IngestionStage,
        percent: float,
        message: str,
        *,
        entities: int = 0,
        relationships: int = 0,
        chunks: int = 0,
    ) -> None:
        progress = IngestionProgress(
            job_id=job_id,
            document_id=document_id,
            filename=filename,
            stage=stage,
            percent=percent,
            message=message,
            entities_found=entities,
            relationships_found=relationships,
            chunks_created=chunks,
        )
        await self._publish(Topic.INGESTION_PROGRESS, **progress.model_dump(mode="json"))
        if callback is not None:
            await callback(progress)

    async def _publish(self, topic: Topic, **payload: object) -> None:
        event = Event.make(topic, source=self.name, **payload)
        await self._deps.events.publish(event.topic.value, event.model_dump(mode="json"))


__all__ = ["IngestionAgent"]
